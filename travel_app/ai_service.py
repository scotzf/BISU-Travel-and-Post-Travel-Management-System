# travel_app/ai_service.py
import os
import re
import json
import requests
import logging
import platform

logger = logging.getLogger(__name__)

OLLAMA_URL     = 'http://localhost:11434/api/generate'
OLLAMA_MODEL   = 'llama3.2:3b'
OLLAMA_TIMEOUT = 60

if platform.system() == 'Windows':
    import pytesseract
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


# ══════════════════════════════════════════════════════════════════════
# DOC TYPE GROUPS
# ══════════════════════════════════════════════════════════════════════

# Full extraction — destination, dates, purpose, traveler names
FULL_EXTRACTION_TYPES = {'TRAVEL_ORDER'}

# Amount only — for budget deduction confirmation
AMOUNT_ONLY_TYPES = {'BURS', 'ITINERARY'}

# No extraction — just store the file
SKIP_EXTRACTION_TYPES = {'DV', 'CERTIFICATE', 'RECEIPTS', 'POST_REPORT', 'LETTER_REQUEST'}


# ══════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════════

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return _extract_from_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        return _extract_from_docx(file_path)
    elif ext in ('.xlsx', '.xls'):
        return _extract_from_xlsx(file_path)
    elif ext in ('.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp'):
        return _extract_from_image(file_path)
    elif ext in ('.txt', '.csv'):
        return _extract_from_text(file_path)
    else:
        return None, 'unsupported'


def _extract_from_pdf(file_path):
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text = ''
        for page in reader.pages:
            text += page.extract_text() or ''
        if len(text.strip()) > 50:
            return text.strip(), 'pdf_text'
        return _ocr_pdf(file_path)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return None, 'error'


def _ocr_pdf(file_path):
    try:
        from pdf2image import convert_from_path
        import pytesseract
        pages = convert_from_path(file_path, dpi=200)
        text  = ''
        for page in pages:
            text += pytesseract.image_to_string(page, lang='eng') + '\n'
        return text.strip(), 'pdf_ocr'
    except Exception as e:
        logger.error(f"PDF OCR error (is poppler-utils installed?): {e}")
        return None, 'error'


def _extract_from_docx(file_path):
    try:
        from docx import Document
        doc  = Document(file_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([c.text.strip() for c in row.cells if c.text.strip()])
                if row_text:
                    text += '\n' + row_text
        return text.strip(), 'docx'
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return None, 'error'


def _extract_from_xlsx(file_path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        ws = wb.active
        text = f'[Sheet: {ws.title}]\n'
        for row in ws.iter_rows(max_row=60, values_only=True):
            row_text = ' | '.join([
                str(v).strip() for v in row
                if v is not None and str(v).strip()
            ])
            if row_text.strip():
                text += row_text + '\n'
        return text.strip(), 'xlsx'
    except Exception as e:
        logger.error(f"XLSX extraction error: {e}")
        return None, 'error'


def _extract_from_image(file_path):
    try:
        import pytesseract
        from PIL import Image
        img  = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang='eng')
        return text.strip(), 'image_ocr'
    except Exception as e:
        logger.error(f"Image OCR error: {e}")
        return None, 'error'


def _extract_from_text(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read().strip(), 'text'
    except Exception as e:
        logger.error(f"Text extraction error: {e}")
        return None, 'error'


# ══════════════════════════════════════════════════════════════════════
# OLLAMA CALLS
# ══════════════════════════════════════════════════════════════════════

def _call_ollama(prompt):
    """Base Ollama call. Returns raw response string or None."""
    try:
        r = requests.post(
            OLLAMA_URL,
            json={'model': OLLAMA_MODEL, 'prompt': prompt, 'stream': False},
            timeout=OLLAMA_TIMEOUT
        )
        if r.status_code != 200:
            logger.error(f"Ollama returned {r.status_code}")
            return None
        return r.json().get('response', '').strip()
    except requests.exceptions.Timeout:
        logger.error("Ollama timeout")
        return None
    except Exception as e:
        logger.error(f"Ollama error: {e}")
        return None


def _parse_json_response(raw):
    if not raw:
        return None
    try:
        if '```' in raw:
            parts = raw.split('```')
            raw = parts[1] if len(parts) > 1 else raw
            if raw.startswith('json'):
                raw = raw[4:]
        start = raw.find('{')
        end   = raw.rfind('}') + 1
        if start >= 0 and end > start:
            return json.loads(raw[start:end])
        # ── Repair: closing brace missing ──
        if start >= 0:
            return json.loads(raw[start:] + '}')
    except json.JSONDecodeError as e:
        logger.error(f"JSON parse error: {e}")
    return None


# ══════════════════════════════════════════════════════════════════════
# TRAVEL ORDER EXTRACTION (full)
# ══════════════════════════════════════════════════════════════════════

def _extract_travel_order(text):
    """
    Extract destination, travel dates, purpose, and traveler names
    from a Travel Order document.
    """
    prompt = f"""You are extracting data from a Philippine government Travel Order document.

Document text:
---
{text[:3000]}
---

Instructions:
1. TRAVELER NAMES — Look for a "To:" section listing names. Extract ONLY the person names, ignore their titles/roles (like "- CTE Faculty", "- Driver", "Chairperson" etc). Include names with titles like "DR." or "PROF." but strip the role descriptions after the dash.
                    - Double check traveler_names — every name listed under "To:" must be in the list, none should be skipped
2. TRAVEL DATES — Look for the actual travel/event dates in the body of the letter (e.g. "October 1-3, 2026", "on October 1 to 3, 2026"). Do NOT use the document date (the date after "Date :"). The travel dates are usually inside the sentence that says "directed to attend" or "hereby authorized".

3. DESTINATION — Look for the venue or city in the body of the letter, usually after "at the" or "in".

4. PURPOSE — The sentence starting with "You are hereby directed to attend..." or "is hereby authorized to travel to...". Keep it concise.

Rules:
- Dates must be in YYYY-MM-DD format
- traveler_names must be a JSON list of full name strings only
- If a date range is given like "November 5-7", start_date is the 5th and end_date is the 7th
- Use null if a field is not found
- Respond with ONLY the JSON object, no explanation

{{
    "destination": "city or venue or null",
    "start_date": "YYYY-MM-DD or null",
    "end_date": "YYYY-MM-DD or null",
    "purpose": "brief purpose or null",
    "traveler_names": ["Full Name", ...],
    "confidence": "high/medium/low"
}}"""

    raw    = _call_ollama(prompt)
    result = _parse_json_response(raw)

    if not result:
        result = _fallback_travel_order(text)
        raw = _call_ollama(prompt)
        print("OLLAMA RAW RESPONSE:", raw)  # check your terminal
        result = _parse_json_response(raw)
        print("PARSED RESULT:", result)

    return result


def _fallback_travel_order(text):
    """Regex fallback if Ollama fails for Travel Order."""
    from datetime import datetime

    result = {'traveler_names': [], 'confidence': 'low'}

    # ── Destination ───────────────────────────────────────────────────
    CITY_PATTERN = (
        r'(Tagbilaran|Cebu|Manila|Davao|Cagayan de Oro|'
        r'Dumaguete|Bacolod|Iloilo|Zamboanga|Bohol|'
        r'Candijay|Bilar|Jagna|Panglao|Ubay|Talibon)'
        r'(?:\s*City)?'
    )
    city = re.search(CITY_PATTERN, text, re.IGNORECASE)
    if city:
        result['destination'] = city.group(0).strip()

    # ── Travel dates (range like "November 5-7, 2025") ────────────────
    range_pattern = (
        r'(January|February|March|April|May|June|July|August|'
        r'September|October|November|December)'
        r'\s+(\d{1,2})-(\d{1,2}),?\s+(\d{4})'
    )
    range_match = re.search(range_pattern, text)
    if range_match:
        month     = range_match.group(1)
        day_start = range_match.group(2)
        day_end   = range_match.group(3)
        year      = range_match.group(4)
        try:
            result['start_date'] = datetime.strptime(
                f"{month} {day_start} {year}", '%B %d %Y'
            ).strftime('%Y-%m-%d')
            result['end_date'] = datetime.strptime(
                f"{month} {day_end} {year}", '%B %d %Y'
            ).strftime('%Y-%m-%d')
        except ValueError:
            pass
    else:
        # Single date pattern — look for it in body text, not document date line
        MONTH_PATTERN = (
            r'(January|February|March|April|May|June|'
            r'July|August|September|October|November|December)'
            r'\s+\d{1,2},?\s+\d{4}'
        )
        # Skip the "Date :" line, find dates in the body
        body_lines = [
            l for l in text.split('\n')
            if not re.match(r'^\s*Date\s*:', l, re.IGNORECASE)
        ]
        body_text = '\n'.join(body_lines)
        dates = list(re.finditer(MONTH_PATTERN, body_text))
        def parse(d):
            for fmt in ['%B %d, %Y', '%B %d %Y']:
                try:
                    return datetime.strptime(d.strip(), fmt).strftime('%Y-%m-%d')
                except ValueError:
                    continue
            return None
        if dates:
            result['start_date'] = parse(dates[0].group(0))
        if len(dates) > 1:
            result['end_date'] = parse(dates[-1].group(0))

    # ── Traveler names (lines after "To :") ───────────────────────────
    lines      = text.split('\n')
    in_to_block = False
    names       = []

    for line in lines:
        stripped = line.strip()

        # Detect start of "To :" block
        if re.match(r'^To\s*:', stripped, re.IGNORECASE):
            in_to_block = True
            # Name may be on same line as "To :"
            name_part = re.sub(r'^To\s*:\s*', '', stripped, flags=re.IGNORECASE).strip()
            name_part = re.sub(r'\s*-\s*.+$', '', name_part).strip()  # strip role
            if name_part:
                names.append(name_part)
            continue

        if in_to_block:
            # Empty line or a line starting a new section ends the block
            if not stripped or re.match(r'^(Date|Subject|Sir|Ma\'am|You are|Your travel)', stripped, re.IGNORECASE):
                in_to_block = False
                continue
            # Strip role descriptions after dash
            name_part = re.sub(r'\s*-\s*.+$', '', stripped).strip()
            if name_part and len(name_part) > 2:
                names.append(name_part)

    if names:
        result['traveler_names'] = names

    # ── Purpose ───────────────────────────────────────────────────────
    for line in lines:
        if len(line) > 40 and any(w in line.lower() for w in [
            'directed to attend', 'authorized to travel',
            'hereby directed', 'hereby authorized'
        ]):
            result.setdefault('purpose', line.strip())
            break

    return result


def _fallback_travel_order(text):
    """Regex fallback if Ollama fails for Travel Order."""
    from datetime import datetime

    MONTH_PATTERN = (
        r'(January|February|March|April|May|June|'
        r'July|August|September|October|November|December)'
        r'\s+\d{1,2},?\s+\d{4}'
    )
    CITY_PATTERN = (
        r'(Tagbilaran|Cebu|Manila|Davao|Cagayan de Oro|'
        r'Dumaguete|Bacolod|Iloilo|Zamboanga|Bohol|'
        r'Candijay|Bilar|Jagna|Panglao|Ubay|Talibon)'
        r'(?:\s*City)?'
    )

    result = {'traveler_names': [], 'confidence': 'low'}

    # Destination
    city = re.search(CITY_PATTERN, text, re.IGNORECASE)
    if city:
        result['destination'] = city.group(0).strip()

    # Dates
    dates = list(re.finditer(MONTH_PATTERN, text))
    def parse(d):
        for fmt in ['%B %d, %Y', '%B %d %Y']:
            try:
                return datetime.strptime(d.strip(), fmt).strftime('%Y-%m-%d')
            except ValueError:
                continue
        return None

    if dates:
        result['start_date'] = parse(dates[0].group(0))
    if len(dates) > 1:
        result['end_date'] = parse(dates[-1].group(0))

    # Purpose — long lines mentioning travel keywords
    for line in text.split('\n'):
        if len(line) > 60 and any(w in line.lower() for w in [
            'travel', 'attend', 'participate', 'training',
            'seminar', 'meeting', 'conference', 'workshop'
        ]):
            result.setdefault('purpose', line.strip())
            break

    return result


# ══════════════════════════════════════════════════════════════════════
# AMOUNT EXTRACTION (BURS / ITINERARY)
# ══════════════════════════════════════════════════════════════════════

def _extract_amount(text, doc_type):
    """
    Extract the total amount from a BURS or Itinerary document.
    """
    prompt = f"""Extract the total amount from this Philippine government {doc_type} document.
Return ONLY a JSON object, no explanation.

Document:
---
{text[:2000]}
---

Rules:
- amount must be numeric only (e.g. 5400.00), no currency symbols
- look for labels like "Total Amount", "Amount Due", "PHP", "₱"
- if not found, use null

{{
    "amount": numeric or null,
    "confidence": "high/medium/low"
}}"""

    raw    = _call_ollama(prompt)
    result = _parse_json_response(raw)

    if not result or result.get('amount') is None:
        result = _fallback_amount(text)

    return result


def _fallback_amount(text):
    """Regex fallback for amount extraction."""
    result = {'amount': None, 'confidence': 'low'}

    for line in text.split('\n'):
        low = line.lower()
        if any(w in low for w in ['total', 'amount due', 'php', '₱']):
            amounts = re.findall(r'\b(\d[\d,]*(?:\.\d{2})?)\b', line)
            # Filter out years and tiny numbers
            amounts = [
                a for a in amounts
                if not re.match(r'^20\d{2}$', a)
                and float(a.replace(',', '')) > 10
            ]
            if amounts:
                result['amount'] = amounts[-1].replace(',', '')
                result['confidence'] = 'low'
                break

    return result


# ══════════════════════════════════════════════════════════════════════
# MAIN ENTRY POINT
# ══════════════════════════════════════════════════════════════════════

def extract_from_document(travel_document):
    """
    Main extraction entry point called after a TravelDocument is saved.

    TRAVEL_ORDER  → full extraction (destination, dates, purpose, traveler_names)
    BURS          → amount only
    ITINERARY     → amount only
    Everything else → skip
    """
    from decimal import Decimal, InvalidOperation
    from datetime import datetime

    doc      = travel_document
    doc_type = doc.doc_type

    # ── Skip types with no extraction ────────────────────────────────
    if doc_type in SKIP_EXTRACTION_TYPES:
        doc.extraction_attempted = False
        doc.save(update_fields=['extraction_attempted'])
        logger.info(f"Doc {doc.id} ({doc_type}) — skipped, no extraction needed")
        return

    # ── Mark as attempted ─────────────────────────────────────────────
    doc.extraction_attempted = True
    doc.save(update_fields=['extraction_attempted'])

    # ── Extract raw text ──────────────────────────────────────────────
    try:
        file_path = doc.file.path
    except Exception:
        _mark_failed(doc, 'Cannot access file path')
        return

    text, method = extract_text_from_file(file_path)

    if not text or len(text.strip()) < 20:
        _mark_failed(doc, f'Could not extract text (method: {method})')
        return

    logger.info(f"Doc {doc.id} ({doc_type}) — text extracted via {method}, {len(text)} chars")

    # ── Route to correct extractor ────────────────────────────────────
    if doc_type in FULL_EXTRACTION_TYPES:
        result = _extract_travel_order(text)
        _save_travel_order_result(doc, result)

    elif doc_type in AMOUNT_ONLY_TYPES:
        result = _extract_amount(text, doc_type)
        _save_amount_result(doc, result)


def _save_travel_order_result(doc, result):
    """Save full Travel Order extraction result to model fields."""
    from datetime import datetime

    if not result:
        _mark_failed(doc, 'Travel Order extraction returned no result')
        return

    update_fields = ['extraction_successful', 'extraction_raw']

    doc.extraction_successful = True
    doc.extraction_raw        = json.dumps(result)

    if result.get('destination'):
        doc.extracted_destination = str(result['destination'])[:200]
        update_fields.append('extracted_destination')

    if result.get('purpose'):
        doc.extracted_purpose = str(result['purpose'])[:500]
        update_fields.append('extracted_purpose')

    for field, model_field in [
        ('start_date', 'extracted_start_date'),
        ('end_date',   'extracted_end_date'),
    ]:
        val = result.get(field)
        if val:
            try:
                setattr(doc, model_field, datetime.strptime(str(val), '%Y-%m-%d').date())
                update_fields.append(model_field)
            except ValueError:
                pass

    doc.save(update_fields=update_fields)
    logger.info(
        f"Doc {doc.id} (TRAVEL_ORDER) — saved. "
        f"Destination: {doc.extracted_destination}, "
        f"Confidence: {result.get('confidence', 'low')}"
    )


def _save_amount_result(doc, result):
    """Save amount extraction result to model fields."""
    from decimal import Decimal, InvalidOperation

    if not result:
        _mark_failed(doc, 'Amount extraction returned no result')
        return

    update_fields = ['extraction_successful', 'extraction_raw']

    doc.extraction_successful = True
    doc.extraction_raw        = json.dumps(result)

    if result.get('amount') is not None:
        try:
            doc.extracted_amount = Decimal(str(result['amount']))
            update_fields.append('extracted_amount')
        except (InvalidOperation, ValueError):
            logger.warning(f"Doc {doc.id} — could not parse amount: {result.get('amount')}")

    doc.save(update_fields=update_fields)
    logger.info(
        f"Doc {doc.id} ({doc.doc_type}) — saved. "
        f"Amount: {doc.extracted_amount}"
    )

def _mark_failed(doc, reason):
    doc.extraction_successful = False
    doc.save(update_fields=['extraction_successful'])
    logger.error(f"Doc {doc.id} extraction failed: {reason}")